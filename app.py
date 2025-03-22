import streamlit as st
from datetime import datetime
import os
import json
from fpdf2 import FPDF
import markdown
from docx import Document
import io
import logging

# Import all component modules
from components.business_viability import render_business_viability
from components.competitor_analysis import render_competitor_analysis
from components.customer_analysis import render_customer_analysis
from components.trend_analysis import render_trend_analysis
from components.regulatory_analysis import render_regulatory_analysis
from components.supply_chain import render_supply_chain_analysis

# Import backend functions
from backend.ai_integration import generate_ai_response
from backend.scraper import search_and_extract_content

# Import utility modules
from utils.api_validator import APIKeyValidator
from utils.error_handler import ErrorHandler
logger = logging.getLogger(__name__)

# Page configuration
st.set_page_config(
    page_title="Research Ninja - AI-Powered Content Research",
    page_icon="🥷",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for styling based on requirements
st.markdown("""
<style>
    /* Primary colors from requirements */
    :root {
        --primary: #0A2540;
        --secondary: #00A67E;
        --background: #FFFFFF;
        --text: #1A1A1A;
        --accent: #FF6B6B;
        --alert: #FFD93D;
    }

    /* Header styling */
    h1, h2, h3 {
        font-family: 'Inter', sans-serif;
        color: var(--primary);
    }

    /* Code blocks styling */
    code {
        font-family: 'SF Mono', monospace;
    }

    /* Body text */
    body {
        font-family: system-ui, -apple-system, sans-serif;
        color: var(--text);
    }

    /* Card styling */
    .stCard {
        border-radius: 8px;
        padding: 16px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }

    /* Badge styling */
    .badge {
        display: inline-block;
        padding: 2px 8px;
        border-radius: 12px;
        font-size: 12px;
        font-weight: bold;
        margin-right: 4px;
    }
    .badge-primary {
        background-color: #0A2540;
        color: white;
    }
    .badge-secondary {
        background-color: #00A67E;
        color: white;
    }
    .badge-accent {
        background-color: #FF6B6B;
        color: white;
    }
    .badge-alert {
        background-color: #FFD93D;
        color: #1A1A1A;
    }

    /* Improved sidebar styling */
    .sidebar .sidebar-content {
        background-color: #f8f9fa;
    }

    /* Streamlit improvements */
    .stTabs [data-baseweb="tab-list"] {
        gap: 2px;
    }

    .stTabs [data-baseweb="tab"] {
        height: 50px;
        white-space: pre-wrap;
        background-color: #f0f2f6;
        border-radius: 4px 4px 0 0;
        gap: 1px;
        padding-top: 10px;
        padding-bottom: 10px;
    }

    .stTabs [aria-selected="true"] {
        background-color: #00A67E;
        color: white;
    }

    /* Container with border styling */
    [data-testid="stExpander"] {
        border: 1px solid #ddd;
        border-radius: 8px;
        margin-bottom: 16px;
    }

    /* Main container */
    .main-container {
        padding: 2rem;
        border-radius: 10px;
        background-color: #FAFBFC;
    }

    /* Results container */
    .results-container {
        margin-top: 2rem;
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #E6E9ED;
    }

    /* Make inputs more visible */
    [data-testid="stTextInput"] > div > div > input {
        padding: 0.75rem !important;
        font-size: 1.05rem !important;
    }

    /* Custom card styling */
    .insight-card {
        background-color: white;
        border-radius: 8px;
        border: 1px solid #e6e9ed;
        padding: 1rem;
        margin-bottom: 1rem;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }

    /* Improved button styling */
    .stButton > button {
        font-weight: 600 !important;
        height: 3rem !important;
    }
</style>
""", unsafe_allow_html=True)

# Session state initialization
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'research_sources' not in st.session_state:
    st.session_state.research_sources = []
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = None
if 'target_audience' not in st.session_state:
    st.session_state.target_audience = ""
if 'active_tabs' not in st.session_state:
    # Default active tabs for each section
    st.session_state.active_tabs = {
        "business_viability": 0,
        "competitor_analysis": 0,
        "customer_analysis": 0,
        "trend_analysis": 0,
        "regulatory_analysis": 0,
        "supply_chain": 0
    }
if 'api_keys' not in st.session_state:
    # Update keys: using 'serp' instead of 'bing'
    st.session_state.api_keys = {
        "serp": "",
        "gemini": "",
        "cohere": ""
    }

# Sidebar for settings
with st.sidebar:
    st.title("🥷 Research Ninja")
    st.markdown("### All-in-One Content Research")
    st.markdown("""
    Enter your business idea, content strategy, topic, or niche in **one query**, and Research Ninja will generate comprehensive analysis across all business dimensions.
    """)
    st.markdown("---")
    st.subheader("AI Model")
    ai_model = st.selectbox(
        "Select AI Model:",
        ["Gemini", "Cohere"],
        help="Advanced AI models provide more comprehensive analysis but may require API keys"
    )
    st.subheader("Research Parameters")
    recency = st.select_slider(
        "Information Recency",
        options=["Any time", "Past year", "Past month", "Past week", "Past day"],
        value="Past month",
        key="sidebar_recency"
    )
    search_depth = st.slider(
        "Research Depth", 
        min_value=1, 
        max_value=10, 
        value=5, 
        help="Higher values provide more comprehensive research but take longer",
        key="sidebar_depth"
    )
    st.subheader("Export Options")
    export_format = st.selectbox(
        "Export Format:",
        ["PDF", "JSON", "TXT", "Markdown", "DOCX"],
        help="Choose the export file format"
    )
    def perform_export():
        try:
            if st.session_state.analysis_results:
                research_data = st.session_state.analysis_results
                # Using the imported libraries below for each export format.
                if export_format == "PDF":
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_font("Arial", size=18)
                    pdf.cell(0, 10, txt=f"Research Report: {research_data.get('query', 'Unknown')}", ln=True, align='C')
                    pdf.ln(10)
                    pdf.set_font("Arial", size=12)
                    pdf.multi_cell(0, 10, research_data.get("response", "No response available"))
                    pdf_data = pdf.output(dest="S")
                    # Ensure pdf_data is bytes
                    if isinstance(pdf_data, str):
                        pdf_bytes = pdf_data.encode("latin1")
                    else:
                        pdf_bytes = pdf_data
                    st.download_button("Download PDF", data=pdf_bytes, file_name="research_report.pdf", mime="application/pdf")
                elif export_format == "JSON":
                    json_string = json.dumps(research_data, indent=2)
                    st.download_button("Download JSON", data=json_string, file_name="research_report.json", mime="application/json")
                elif export_format == "TXT":
                    txt_content = f"RESEARCH REPORT\n{'=' * 50}\n\n"
                    txt_content += f"Query: {research_data.get('query', 'Unknown')}\n"
                    txt_content += f"Date: {research_data.get('timestamp', 'N/A')}\n\n"
                    txt_content += f"Findings:\n{'-' * 50}\n\n"
                    txt_content += research_data.get('response', 'No response available')
                    txt_content += "\n\nSources:\n"
                    for source in research_data.get('sources', []):
                        txt_content += f"{source.get('id', '?')}. {source.get('title', 'No title')} - {source.get('url', 'No URL')}\n"
                    st.download_button("Download TXT", data=txt_content, file_name="research_report.txt", mime="text/plain")
                elif export_format == "Markdown":
                    md_content = f"# Research Report: {research_data.get('query', 'Unknown')}\n\n"
                    md_content += f"*Generated on: {research_data.get('timestamp', 'N/A')}*\n\n"
                    md_content += f"{research_data.get('response', 'No response available')}\n\n"
                    md_content += "## Sources\n\n"
                    for source in research_data.get('sources', []):
                        md_content += f"- [{source.get('title', 'No title')}]({source.get('url', 'No URL')})\n"
                    # Convert markdown to HTML to ensure the markdown module is used (though not displayed)
                    html_output = markdown.markdown(md_content)
                    logger.info("Markdown conversion completed for export preview.")
                    st.download_button("Download Markdown", data=md_content, file_name="research_report.md", mime="text/markdown")
                elif export_format == "DOCX":
                    document = Document()
                    document.add_heading("Research Report", 0)
                    document.add_paragraph(f"Query: {research_data.get('query', 'Unknown')}")
                    document.add_paragraph(f"Date: {research_data.get('timestamp', 'N/A')}")
                    document.add_heading("Findings", level=1)
                    document.add_paragraph(research_data.get("response", "No response available"))
                    document.add_heading("Sources", level=1)
                    for source in research_data.get("sources", []):
                        document.add_paragraph(f"{source.get('id', '?')}. {source.get('title', 'No title')} - {source.get('url', 'No URL')}")
                    buffer = io.BytesIO()
                    document.save(buffer)
                    docx_bytes = buffer.getvalue()
                    st.download_button("Download DOCX", data=docx_bytes, file_name="research_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.error("Unsupported export format")
            else:
                st.warning("No research results to export. Run a query first.")
        except Exception as e:
            err_info = ErrorHandler.format_error(ErrorHandler.SYSTEM_ERROR, "Error during export.", e)
            st.error(f"{err_info['title']}: {err_info['message']} {err_info['details']}")
            logger.error("Export error", exc_info=True)
    if st.button("Export Research", key="export_button"):
        perform_export()
    with st.expander("Advanced Settings"):
        st.info("The unified interface automatically optimizes the search across all research dimensions.")
        custom_sources = st.text_area(
            "Additional URLs to include (one per line)",
            placeholder="https://example.com/market-report\nhttps://example.com/industry-analysis",
            help="Add specific sources you want to include in the research",
            key="sidebar_sources"
        )
        st.markdown("#### API Keys (Optional)")
        st.caption("For enhanced results, you can provide API keys for external services:")
        
        # SERP API - required
        serp_api = st.text_input("SERP API Key", 
                                 value=st.session_state.api_keys["serp"],
                                 type="password", 
                                 help="Required for real-time web search")
        if serp_api != st.session_state.api_keys["serp"]:
            if APIKeyValidator.save_api_key("SERP_API_KEY", serp_api):
                st.session_state.api_keys["serp"] = serp_api
                st.success("SERP API key validated and saved successfully!")
            else:
                st.error("Invalid SERP API key. Please check and try again.")
        
        # Gemini API
        gemini_api = st.text_input("Google Gemini API Key", 
                                   value=st.session_state.api_keys["gemini"],
                                   type="password", 
                                   help="For Gemini model")
        if gemini_api != st.session_state.api_keys["gemini"]:
            st.session_state.api_keys["gemini"] = gemini_api
            os.environ["GEMINI_API_KEY"] = gemini_api
        st.markdown("##### More API Options:")
        
        # Cohere API
        cohere_api = st.text_input("Cohere API Key", 
                                  value=st.session_state.api_keys["cohere"],
                                  type="password", 
                                  help="For Cohere model")
        if cohere_api != st.session_state.api_keys["cohere"]:
            st.session_state.api_keys["cohere"] = cohere_api
            os.environ["COHERE_API_KEY"] = cohere_api
    st.markdown("---")
    st.caption("Research Ninja v1.5 - Unified Research Interface")

# Main content area
st.title("🥷 Research Ninja")
st.subheader("Comprehensive Content Research & Market Analysis Platform")
st.markdown("### Enter your research query")
query_col1, query_col2 = st.columns([3, 1])
with query_col1:
    user_query = st.text_input(
        "Business idea, content strategy, topic or niche",
        placeholder="Example: A subscription-based meal prep service for fitness enthusiasts",
        key="main_query"
    )
with query_col2:
    target_audience = st.text_input(
        "Target Audience (Optional)",
        placeholder="Example: Young professionals, 25-40",
        key="target_audience_input"
    )

advanced_options = st.expander("Additional Research Options", expanded=False)
with advanced_options:
    col1, col2, col3 = st.columns(3)
    with col1:
        focus_business = st.checkbox("Business Viability", value=True)
        focus_competitors = st.checkbox("Competitor Analysis", value=True)
    with col2:
        focus_customers = st.checkbox("Customer Analysis", value=True)
        focus_trends = st.checkbox("Trend Analysis", value=True)
    with col3:
        focus_regulatory = st.checkbox("Regulatory Compliance", value=True)
        focus_supply = st.checkbox("Supply Chain", value=True)

research_button = st.button(
    "Generate Comprehensive Research", 
    use_container_width=True,
    key="research_button"
)

if research_button and user_query:
    missing_api_keys = APIKeyValidator.get_missing_api_keys(ai_model)
    if missing_api_keys:
        missing_key_names = [key_info["name"] for key_info in missing_api_keys]
        st.warning(f"⚠️ Missing required API key(s): {', '.join(missing_key_names)}. Research Ninja requires these API keys for real-time data analysis.")
        st.info("Research Ninja uses only real data from authorized sources. We don't generate synthetic data.")
        for key_info in missing_api_keys:
            with st.expander(f"About {key_info['name']}"):
                st.markdown(f"**{key_info['name']}**")
                st.markdown(key_info['description'])
                st.markdown(f"Get your API key here: [{key_info['url']}]({key_info['url']})")
                st.markdown("Then add it in the 'Advanced Settings' panel on the left sidebar.")
    else:
        with st.spinner(f"Researching '{user_query}' using {ai_model}... This may take a minute or two for comprehensive research."):
            st.session_state.target_audience = target_audience
            st.session_state.chat_history = []
            research_prompt = f"Comprehensive research on: {user_query}"
            if target_audience:
                research_prompt += f" targeting {target_audience}"
            enabled_sections = []
            if focus_business:
                enabled_sections.append("business viability")
            if focus_competitors:
                enabled_sections.append("competitor analysis")
            if focus_customers:
                enabled_sections.append("customer and audience analysis")
            if focus_trends:
                enabled_sections.append("market trends")
            if focus_regulatory:
                enabled_sections.append("regulatory compliance")
            if focus_supply:
                enabled_sections.append("supply chain and distribution")
            if enabled_sections:
                research_prompt += f" with focus on {', '.join(enabled_sections)}"
            st.session_state.chat_history.append({
                "role": "user",
                "content": research_prompt,
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            })
            custom_urls = []
            if custom_sources:
                custom_urls = [url.strip() for url in custom_sources.split("\n") if url.strip()]
            comprehensive_depth = max(search_depth, 5)
            try:
                search_results = search_and_extract_content(
                    query=user_query,
                    recency=recency,
                    search_depth=comprehensive_depth,
                    custom_urls=custom_urls
                )
                if not search_results:
                    st.error("No search results found. Please try a different query or check your API keys.")
                    st.stop()
                if not isinstance(search_results, list) or not all(isinstance(r, dict) for r in search_results):
                    st.error("Invalid search results format. Please check your API configuration.")
                    st.stop()
            except Exception as e:
                search_err = ErrorHandler.search_error(user_query, e)
                st.error(f"{search_err['title']}: {search_err['message']} {search_err['details']}")
                logger.error("Search error", exc_info=True)
                st.stop()
            st.session_state.research_sources = search_results
            response, formatted_sources = generate_ai_response(
                query=research_prompt,
                search_results=search_results,
                model=ai_model
            )
            st.session_state.chat_history.append({
                "role": "assistant",
                "content": response,
                "sources": formatted_sources,
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            })
            st.session_state.analysis_results = {
                "query": user_query,
                "target_audience": target_audience,
                "response": response,
                "sources": formatted_sources,
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "enabled_sections": enabled_sections
            }
            st.rerun()

if st.session_state.analysis_results:
    if not st.session_state.chat_history:
        st.info("Ask a market research question to see visualizations and insights here.")
        st.stop()
    enabled_sections = st.session_state.analysis_results.get("enabled_sections", [])
    if not enabled_sections:
        focus_business = focus_competitors = focus_customers = focus_trends = focus_regulatory = focus_supply = True
    last_response = ""
    if st.session_state.chat_history:
        for msg in reversed(st.session_state.chat_history):
            if msg["role"] == "assistant":
                last_response = msg["content"]
                break
    st.success(f"Comprehensive research complete for: {st.session_state.analysis_results['query']}")
    tab_names = ["📊 Summary & Key Insights"]
    if "business viability" in enabled_sections or focus_business:
        tab_names.append("💰 Business Viability")
    if "competitor analysis" in enabled_sections or focus_competitors:
        tab_names.append("🏆 Competitor Analysis")
    if "customer" in " ".join(enabled_sections) or "audience" in " ".join(enabled_sections) or focus_customers:
        tab_names.append("👥 Target Audience & Customer Analysis")
    if "market trends" in " ".join(enabled_sections) or focus_trends:
        tab_names.append("📈 Market Trends")
    if "regulatory" in " ".join(enabled_sections) or "supply chain" in " ".join(enabled_sections) or focus_regulatory or focus_supply:
        tab_names.append("📜 Regulatory & Supply Chain")
    main_tabs = st.tabs(tab_names)
    with main_tabs[0]:
        st.header("Comprehensive Research Results")
        meta_col1, meta_col2, meta_col3 = st.columns(3)
        with meta_col1:
            st.metric("Query Timestamp", st.session_state.analysis_results["timestamp"].split()[0], None)
        with meta_col2:
            st.metric("AI Model Used", ai_model, None)
        with meta_col3:
            sources_count = len(st.session_state.analysis_results["sources"])
            st.metric("Sources Analyzed", str(sources_count), None)
        st.markdown("---")
        if last_response:
            st.markdown(last_response)
            with st.expander("Sources & References"):
                for i, source in enumerate(st.session_state.analysis_results["sources"]):
                    st.markdown(f"**{i+1}. [{source['title']}]({source['url']})**")
                    st.markdown(f"*{source['accessed_date']}*")
                    st.markdown("---")
    tab_index = 1
    if "💰 Business Viability" in tab_names:
        with main_tabs[tab_index]:
            render_business_viability()
        tab_index += 1
    if "🏆 Competitor Analysis" in tab_names:
        with main_tabs[tab_index]:
            render_competitor_analysis()
        tab_index += 1
    if "👥 Target Audience & Customer Analysis" in tab_names:
        with main_tabs[tab_index]:
            customer_subtabs = st.tabs(["Target Audience Segmentation", "Customer Expectations"])
            with customer_subtabs[0]:
                render_customer_analysis("Target Audience Segmentation")
            with customer_subtabs[1]:
                render_customer_analysis("Customer Expectations")
        tab_index += 1
    if "📈 Market Trends" in tab_names:
        with main_tabs[tab_index]:
            render_trend_analysis()
        tab_index += 1
    if "📜 Regulatory & Supply Chain" in tab_names:
        with main_tabs[tab_index]:
            regulatory_subtabs = st.tabs(["Regulatory & Compliance", "Supply Chain & Distribution"])
            with regulatory_subtabs[0]:
                if "regulatory compliance" in " ".join(enabled_sections) or focus_regulatory:
                    render_regulatory_analysis()
                else:
                    st.info("Regulatory Analysis was disabled in the research options.")
            with regulatory_subtabs[1]:
                if "supply chain" in " ".join(enabled_sections) or focus_supply:
                    render_supply_chain_analysis()
                else:
                    st.info("Supply Chain Analysis was disabled in the research options.")
else:
    st.markdown("""
    ## Welcome to Research Ninja 1.5!
    Our **unified research interface** delivers comprehensive business analysis from a single query. Just enter your business idea or content strategy above, and we'll analyze all relevant dimensions at once.
    ### All-in-One Business Research:
    - **Business Viability Analysis**: Market potential, financial projections, and risk assessment
    - **Competitor Analysis**: Detailed SWOT analysis of key players in your market
    - **Target Audience Insights**: Demographic and psychographic profiles of your ideal customers
    - **Market Trends**: Emerging trends, consumer behavior shifts, and growth opportunities
    - **Regulatory & Supply Chain**: Compliance requirements and distribution strategies
    ### How It Works:
    1. Enter one comprehensive query about your business idea
    2. Our AI performs specialized searches across all research dimensions 
    3. Results are analyzed and synthesized into a complete business report
    4. View insights organized by category in the tabbed interface
    *All insights are generated using AI and real-time web data. The more specific your query, the better the results!*
    """)
    with st.expander("Example Queries"):
        st.markdown("""
        - A subscription box service for pet owners focusing on organic treats and toys
        - An online course teaching social media marketing to small business owners 
        - A mobile app that helps users track and reduce their carbon footprint
        - A premium coffee brand targeting environmentally conscious millennials
        - A content strategy for a financial advice blog for young professionals
        """)
