import base64
import json
import io
import logging
from typing import Dict, Any
from datetime import datetime

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def export_to_pdf(research_data: Dict[str, Any]) -> str:
    """
    Export research data to a PDF document and return it as a base64 encoded string.
    
    Uses the reportlab library to generate the PDF.
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

        # Use a context manager for the buffer
        with io.BytesIO() as buffer:
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()

            # Define custom styles
            styles.add(ParagraphStyle(
                name='CustomTitle',
                fontSize=18,
                spaceAfter=12,
                textColor=colors.HexColor('#0A2540')
            ))
            styles.add(ParagraphStyle(
                name='CustomSubtitle',
                fontSize=14,
                spaceAfter=8,
                textColor=colors.HexColor('#00A67E')
            ))
            # Overwrite the Normal style to add spacing if needed
            styles['Normal'].spaceAfter = 6
            styles['Normal'].fontSize = 10

            content = []
            content.append(Paragraph(
                f"Research Report: {research_data.get('query', 'Unknown')}",
                styles['CustomTitle']
            ))
            content.append(Spacer(1, 12))
            content.append(Paragraph(
                f"Date: {research_data.get('timestamp', 'N/A')}",
                styles['Normal']
            ))
            content.append(Spacer(1, 20))

            # Split the response into lines and format based on header markers
            response_text = research_data.get('response', '')
            for line in response_text.split('\n'):
                stripped_line = line.strip()
                if stripped_line.startswith('# '):
                    content.append(Paragraph(stripped_line.replace('# ', ''), styles['CustomTitle']))
                elif stripped_line.startswith('## '):
                    content.append(Paragraph(stripped_line.replace('## ', ''), styles['CustomSubtitle']))
                else:
                    content.append(Paragraph(stripped_line, styles['Normal']))
                content.append(Spacer(1, 6))

            doc.build(content)
            pdf_content = buffer.getvalue()

        return base64.b64encode(pdf_content).decode('utf-8')
    except Exception as e:
        logger.exception("Error exporting to PDF")
        return ""


def export_to_json(research_data: Dict[str, Any]) -> str:
    """
    Export research data to JSON format with pretty-printing.
    """
    try:
        return json.dumps(research_data, indent=2)
    except Exception as e:
        logger.exception("Error exporting to JSON")
        return "{}"


def export_to_txt(research_data: Dict[str, Any]) -> str:
    """
    Export research data to plain text format.
    """
    try:
        with io.StringIO() as buffer:
            buffer.write("RESEARCH REPORT\n" + "=" * 50 + "\n\n")
            buffer.write(f"Query: {research_data.get('query', 'Unknown')}\n")
            buffer.write(f"Date: {research_data.get('timestamp', 'N/A')}\n\n")
            buffer.write("Findings:\n" + "-" * 50 + "\n\n")
            buffer.write(research_data.get('response', 'No response available'))
            buffer.write("\n\nSources:\n")
            for source in research_data.get('sources', []):
                buffer.write(
                    f"{source.get('id', '?')}. {source.get('title', 'No title')} - {source.get('url', 'No URL')}\n"
                )
            return buffer.getvalue()
    except Exception as e:
        logger.exception("Error exporting to TXT")
        return ""


def export_to_markdown(research_data: Dict[str, Any]) -> str:
    """
    Export research data to Markdown format.
    """
    try:
        with io.StringIO() as buffer:
            buffer.write(f"# Research Report: {research_data.get('query', 'Unknown')}\n\n")
            buffer.write(f"*Generated on: {research_data.get('timestamp', 'N/A')}*\n\n")
            buffer.write("---\n\n")
            buffer.write(f"{research_data.get('response', 'No response available')}\n\n")
            buffer.write("## Sources\n\n")
            for source in research_data.get('sources', []):
                buffer.write(f"- [{source.get('title', 'No title')}]({source.get('url', 'No URL')})\n")
            return buffer.getvalue()
    except Exception as e:
        logger.exception("Error exporting to Markdown")
        return ""


def export_to_docx(research_data: Dict[str, Any]) -> str:
    """
    Export research data to DOCX format and return it as a base64 encoded string.
    
    Uses the python-docx library to construct the document.
    """
    try:
        from docx import Document
        document = Document()
        document.add_heading("Research Report", 0)
        document.add_paragraph(f"Query: {research_data.get('query', 'Unknown')}")
        document.add_paragraph(f"Date: {research_data.get('timestamp', 'N/A')}")
        document.add_paragraph("\n---\n")
        document.add_paragraph(research_data.get('response', 'No response available'))

        document.add_heading("Sources", level=1)
        for source in research_data.get('sources', []):
            document.add_paragraph(
                f"{source.get('id', '?')}. {source.get('title', 'No title')} - {source.get('url', 'No URL')}"
            )

        with io.BytesIO() as buffer:
            document.save(buffer)
            docx_content = buffer.getvalue()

        return base64.b64encode(docx_content).decode('utf-8')
    except Exception as e:
        logger.exception("Error exporting to DOCX")
        return ""